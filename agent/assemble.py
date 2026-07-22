from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from docx.shared import Inches
import os
import random
from typing import Dict, List, Any
from langchain_core.messages import AIMessage, SystemMessage
from pathlib import Path

# 中文字号映射表
CHINESE_TO_PTS = {
    "初号": 42, "一号": 26, "二号": 22, "三号": 16,
    "四号": 14, "小四": 12, "五号": 10.5, "小五": 9,
}

# 对齐方式映射
ALIGN_MAP = {
    "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
    "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
    "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
    "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
}

# 字体映射
FONT_MAP = {
    "宋体": "宋体",
    "黑体": "黑体",
    "楷体": "楷体",
    "仿宋": "仿宋",
    "times new roman": "Times New Roman",
    "times": "Times New Roman",
}

# 文档组装工具集
DOCUMENT_TOOLS = {
    # 字体格式工具
    "set_font_name", "set_font_size", "set_font_bold",
    "set_font_italic", "set_font_underline",

    # 段落格式工具
    "set_paragraph_align", "set_line_spacing",
    "set_space_before", "set_space_after", "set_paragraph_indent",

    # 文档结构工具
    "add_section_break", "set_page_margins", "add_page_number"
}

# 所有子图Agent的名称列表
ALL_AGENT_NAMES = ["cover", "text", "author", "chart", "fund", "reference", "other"]


def resolve_font_size(size: str) -> float:
    """将中文或字符串转为 Pt 值"""
    if isinstance(size, (int, float)):
        return float(size)
    return CHINESE_TO_PTS.get(size, 12)  # 默认小四


def apply_document_tool(tool_name: str, args: dict, doc, p, run):
    """应用文档工具到文档、段落或运行文本"""
    try:
        if tool_name == "set_font_name":
            font_name = args.get("font_name")
            if font_name:
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

        elif tool_name == "set_font_size":
            pt_val = args.get("pt")
            size_pt = resolve_font_size(pt_val)
            run.font.size = Pt(size_pt)

        elif tool_name == "set_font_bold":
            run.bold = bool(args.get("bold", False))

        elif tool_name == "set_font_italic":
            run.italic = bool(args.get("italic", False))

        elif tool_name == "set_font_underline":
            run.underline = bool(args.get("underline", False))

        elif tool_name == "set_paragraph_align":
            align = args.get("align")
            if align in ALIGN_MAP:
                p.alignment = ALIGN_MAP[align]

        elif tool_name == "set_line_spacing":
            p.paragraph_format.line_spacing = float(args.get("multiple", 1.0))

        elif tool_name == "set_space_before":
            p.paragraph_format.space_before = Pt(float(args.get("pt", 0)))

        elif tool_name == "set_space_after":
            p.paragraph_format.space_after = Pt(float(args.get("pt", 0)))

        elif tool_name == "set_paragraph_indent":
            indent_type = args.get("type", "first_line")
            characters = args.get("characters", 0)
            pt_per_char = 10.5 / 2
            indent_pt = characters * pt_per_char

            if indent_type == "first_line":
                p.paragraph_format.first_line_indent = Pt(indent_pt)
            elif indent_type == "hanging":
                p.paragraph_format.left = Pt(indent_pt)
                p.paragraph_format.first_line_indent = Pt(-indent_pt)

        elif tool_name == "add_section_break":
            doc.add_section(WD_SECTION.NEW_PAGE)

        elif tool_name == "set_page_margins":
            section = doc.sections[-1]
            for margin in ["top", "bottom", "left", "right"]:
                if margin in args:
                    setattr(section, f"{margin}_margin", Inches(float(args[margin])))

        elif tool_name == "add_page_number":
            section = doc.sections[-1]
            footer = section.footer
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.add_run(f"第 {args.get('page', '')} 页")

    except Exception as e:
        raise RuntimeError(f"工具 {tool_name} 应用失败: {e}")


def add_issues_section(doc, verification_issues):
    """在文档末尾添加问题汇总部分"""
    if not verification_issues:
        return

    # 添加分节符
    doc.add_section(WD_SECTION.NEW_PAGE)

    # 添加标题
    p_title = doc.add_paragraph()
    run_title = p_title.add_run("待改进问题汇总")
    run_title.font.size = Pt(16)
    run_title.bold = True
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 添加说明
    p_desc = doc.add_paragraph()
    p_desc.add_run("以下是在文档检查过程中发现的需要改进的问题：")

    # 添加问题列表
    for i, issue in enumerate(verification_issues, 1):
        p_issue = doc.add_paragraph()
        p_issue.paragraph_format.left_indent = Pt(20)
        p_issue.paragraph_format.space_after = Pt(6)

        run_issue = p_issue.add_run(f"{i}. 【{issue.get('section', '未知')}-{issue.get('subsection', '未知')}】")
        run_issue.bold = True

        p_desc2 = doc.add_paragraph()
        p_desc2.paragraph_format.left_indent = Pt(40)
        p_desc2.add_run(f"问题类型: {issue.get('issue_type', '未知')}")

        p_desc3 = doc.add_paragraph()
        p_desc3.paragraph_format.left_indent = Pt(40)
        p_desc3.add_run(f"问题描述: {issue.get('description', '无描述')}")

        p_desc4 = doc.add_paragraph()
        p_desc4.paragraph_format.left_indent = Pt(40)
        run_suggestion = p_desc4.add_run(f"改进建议: {issue.get('suggestion', '无建议')}")
        run_suggestion.font.color.rgb = RGBColor(0, 128, 0)  # 绿色


def collect_all_layout_items(state: Dict[str, Any]) -> List[Dict[str, Any]]:
    """遍历所有Agent结果，收集排版条目并按原始顺序排列"""
    all_items = []

    for agent_name in ALL_AGENT_NAMES:
        agent_state = state.get(agent_name, {})
        agent_missions = agent_state.get("agent_mission", [])
        agent_result = agent_state.get("agent_result", {})
        layout_results = agent_result.get("layout_result", [])

        if not agent_missions or not layout_results:
            continue

        # 构建 mission 查找表
        mission_map = {}
        for mission in agent_missions:
            key = (mission.get("section"), mission.get("subsection"))
            mission_map[key] = mission

        for item in layout_results:
            section = item.get("section")
            subsection = item.get("subsection")
            lookup_key = (section, subsection)

            mission = mission_map.get(lookup_key, {})
            content = mission.get("content", "")

            all_items.append({
                "agent_name": agent_name,
                "section": section,
                "subsection": subsection,
                "content": content,
                "status": item.get("status"),
                "tool_execution_results": item.get("tool_execution_results", []),
            })

    return all_items


def apply_header_footer(doc, all_layout_items, assembly_log):
    """从排版结果中提取并应用页眉页脚页码设置"""
    header_config = None
    footer_config = None
    page_number_config = None

    for item in all_layout_items:
        for exec_res in item.get("tool_execution_results", []):
            if exec_res.get("execution_result", {}).get("status") != "success":
                continue
            action = exec_res.get("execution_result", {}).get("result", {}).get("action", "")
            data = exec_res.get("execution_result", {}).get("result", {}).get("data", {})
            if action == "set_header":
                header_config = data
                assembly_log.append(f"📋 页眉: {data.get('text', '')}")
            elif action == "set_footer":
                footer_config = data
                assembly_log.append(f"📋 页脚: {data.get('text', '')}")
            elif action == "set_page_number":
                page_number_config = data
                assembly_log.append(f"📋 页码: 位置={data.get('position','center_bottom')} 起始={data.get('start_from',1)}")

    if not header_config and not footer_config and not page_number_config:
        return

    SIZE_MAP = {"小五": 9, "五号": 10.5, "小四": 12, "四号": 14, "三号": 16}
    for section in doc.sections:
        # 页眉
        if header_config:
            header = section.header
            header.is_linked_to_previous = False
            hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            hp.text = ""
            run = hp.add_run(header_config.get("text", ""))
            run.font.name = header_config.get("font_name", "宋体")
            run._element.rPr.rFonts.set(qn('w:eastAsia'), header_config.get("font_name", "宋体"))
            size_str = header_config.get("font_size", "小五")
            run.font.size = Pt(SIZE_MAP.get(size_str, 9))
            align_map = {"left": 0, "center": 1, "right": 2}
            hp.alignment = align_map.get(header_config.get("align", "center"), 1)

        # 页脚
        if footer_config:
            footer = section.footer
            footer.is_linked_to_previous = False
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            fp.text = ""
            run = fp.add_run(footer_config.get("text", ""))
            run.font.name = footer_config.get("font_name", "宋体")
            run._element.rPr.rFonts.set(qn('w:eastAsia'), footer_config.get("font_name", "宋体"))
            size_str = footer_config.get("font_size", "小五")
            run.font.size = Pt(SIZE_MAP.get(size_str, 9))
            align_map = {"left": 0, "center": 1, "right": 2}
            fp.alignment = align_map.get(footer_config.get("align", "center"), 1)

        # 页码
        if page_number_config:
            footer = section.footer
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            # 添加页码域
            from docx.oxml import OxmlElement
            fld_char_begin = OxmlElement('w:fldChar')
            fld_char_begin.set(qn('w:fldCharType'), 'begin')
            instr_text = OxmlElement('w:instrText')
            instr_text.set(qn('xml:space'), 'preserve')
            start_from = page_number_config.get("start_from", 1)
            instr_text.text = f'PAGE \\* MERGEFORMAT'
            fld_char_end = OxmlElement('w:fldChar')
            fld_char_end.set(qn('w:fldCharType'), 'end')
            run_page = fp.add_run()
            run_page._r.append(fld_char_begin)
            run_page._r.append(instr_text)
            run_page._r.append(fld_char_end)
            run_page.font.size = Pt(9)
            fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # 设置起始页码
            sect_pr = section._sectPr
            if sect_pr is None:
                sect_pr = OxmlElement('w:sectPr')
                section._element.append(sect_pr)
            pg_num_type = OxmlElement('w:pgNumType')
            pg_num_type.set(qn('w:start'), str(start_from))
            sect_pr.append(pg_num_type)


def assemble_node(state: Dict[str, Any]) -> Dict[str, Any]:
    """文档组装车间 - 将所有Agent的排版结果组装成完整文档"""
    system_messages = [
        SystemMessage(content="<-- 进入 09 文档组装车间 -->")
    ]

    verification_issues = state.get("verification_issues", [])

    # 收集所有Agent的排版结果
    all_layout_items = collect_all_layout_items(state)

    if not all_layout_items:
        system_messages.append(SystemMessage(content="⚠️ 无排版任务结果，生成空文档"))
        doc = Document()
        # 使用项目相对路径
        output_dir = str(Path(__file__).resolve().parent.parent / "data" / "output")
        os.makedirs(output_dir, exist_ok=True)
        random_suffix = random.randint(1000, 9999)
        filename = f"文档组装结果_{random_suffix}.docx"
        output_file = os.path.join(output_dir, filename)
        doc.save(output_file)
        return {
            "messages": system_messages + [SystemMessage(content=f"⚠️ 已生成空文档: {output_file}")],
            "final_paper": os.path.abspath(output_file),
            "awaiting": "user_email",
        }

    # 创建 Word 文档
    doc = Document()

    # 设置默认页面边距
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # 组装文档主体内容（按原始顺序）
    assembly_log = []
    agent_statistics = {}

    # 应用页眉页脚页码
    apply_header_footer(doc, all_layout_items, assembly_log)

    for idx, item in enumerate(all_layout_items):
        content = item.get("content", "")
        agent_name = item.get("agent_name", "unknown")

        if not content:
            assembly_log.append(f"⏭️ [{agent_name}] {item.get('section')}-{item.get('subsection')}: 无内容，跳过")
            continue

        # 如果换了 agent，添加分节标识（第一个不加）
        if idx > 0 and agent_name != all_layout_items[idx - 1].get("agent_name"):
            p_sep = doc.add_paragraph()
            p_sep.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run_sep = p_sep.add_run(f"———— {agent_name} 排版区域 ————")
            run_sep.font.size = Pt(9)
            run_sep.font.color.rgb = RGBColor(180, 180, 180)

        # 添加段落
        p = doc.add_paragraph()
        run = p.add_run(content)

        # 应用所有格式工具
        tool_count = 0
        for exec_res in item.get("tool_execution_results", []):
            if exec_res.get("execution_result", {}).get("status") != "success":
                continue

            tool_name = exec_res.get("tool_name")
            if tool_name in DOCUMENT_TOOLS:
                try:
                    apply_document_tool(
                        tool_name,
                        exec_res.get("args", {}),
                        doc, p, run
                    )
                    tool_count += 1
                except Exception as e:
                    assembly_log.append(f"❌ [{agent_name}] {item.get('subsection')} - {tool_name}: {e}")

        assembly_log.append(f"✅ [{agent_name}] {item.get('subsection') or item.get('section')}: 应用 {tool_count} 个工具")

        # 统计
        if agent_name not in agent_statistics:
            agent_statistics[agent_name] = {"count": 0, "tool_count": 0}
        agent_statistics[agent_name]["count"] += 1
        agent_statistics[agent_name]["tool_count"] += tool_count

    # 添加验证问题汇总
    if verification_issues:
        add_issues_section(doc, verification_issues)
        assembly_log.append(f"📝 添加了 {len(verification_issues)} 个待改进问题")

    # 保存文档（使用项目相对路径）
    output_dir = str(Path(__file__).resolve().parent.parent / "data" / "output")
    os.makedirs(output_dir, exist_ok=True)

    random_suffix = random.randint(1000, 9999)
    filename = f"文档组装结果_{random_suffix}.docx"
    output_file = os.path.join(output_dir, filename)

    try:
        doc.save(output_file)
        abs_path = os.path.abspath(output_file)

        # 添加统计信息到日志
        stats_lines = ["📊 组装统计："]
        for agent, stats in agent_statistics.items():
            stats_lines.append(f"  - {agent}: {stats['count']} 项，{stats['tool_count']} 个工具")

        for log in assembly_log:
            system_messages.append(SystemMessage(content=log))
        for stats_line in stats_lines:
            system_messages.append(SystemMessage(content=stats_line))

        system_messages.append(SystemMessage(content=f"🎉 文档组装完成: {abs_path}"))

        return {
            "messages": system_messages,
            "final_paper": abs_path,
            "awaiting": "user_email",
            "assembly_log": assembly_log,
            "issues_count": len(verification_issues),
            "agent_statistics": agent_statistics,
        }

    except Exception as e:
        system_messages.append(SystemMessage(content=f"❌ 文档保存失败: {str(e)}"))
        return {
            "messages": system_messages,
            "error": f"文档保存失败: {str(e)}",
            "awaiting": ""
        }
